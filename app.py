from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, current_user, login_required
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta
from functools import wraps
from secrets import token_urlsafe
import os
import io
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from PIL import Image
import base64
import re
from docx.text.paragraph import Paragraph



# Initialize Flask app
app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///findings.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = os.urandom(24)
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=8)

# Initialize database
db = SQLAlchemy(app)

# Initialize Login Manager
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login_page'

# User model with UserMixin for Flask-Login
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=True)
    role = db.Column(db.String(20), default='user')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_login = db.Column(db.DateTime)
    reset_token = db.Column(db.String(100), unique=True, nullable=True)
    reset_token_expires = db.Column(db.DateTime, nullable=True)
    keyword1 = db.Column(db.String(100), nullable=True)
    keyword2 = db.Column(db.String(100), nullable=True)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    def set_keywords(self, keyword1, keyword2):
        self.keyword1 = generate_password_hash(keyword1)
        self.keyword2 = generate_password_hash(keyword2)

    def check_keywords(self, keyword1, keyword2):
        return (check_password_hash(self.keyword1, keyword1) and 
                check_password_hash(self.keyword2, keyword2))

    def get_id(self):
        return str(self.id)

# Finding model
class Finding(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    risk_rating = db.Column(db.String(50), nullable=False)
    description = db.Column(db.Text, nullable=False)
    impact = db.Column(db.Text, nullable=False)
    resolution = db.Column(db.Text, nullable=False)
    category = db.Column(db.String(100), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def to_dict(self):
        return {
            'id': self.id,
            'title': self.title,
            'risk_rating': self.risk_rating,
            'description': self.description,
            'impact': self.impact,
            'resolution': self.resolution,
            'category': self.category,
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M:%S')
        }

# UserFinding model
class UserFinding(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    finding_id = db.Column(db.Integer, db.ForeignKey('finding.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    risk_rating = db.Column(db.String(50), nullable=False)
    description = db.Column(db.Text, nullable=False)
    impact = db.Column(db.Text, nullable=False)
    resolution = db.Column(db.Text, nullable=False)
    category = db.Column(db.String(100), nullable=True)
    resources_affected = db.Column(db.Text, nullable=True)
    evidence = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    user = db.relationship('User', backref=db.backref('user_findings', lazy=True))
    finding = db.relationship('Finding')
    
    def to_dict(self):
        return {
            'id': self.id,
            'user_id': self.user_id,
            'finding_id': self.finding_id,
            'title': self.title,
            'risk_rating': self.risk_rating,
            'description': self.description,
            'impact': self.impact,
            'resolution': self.resolution,
            'category': self.category,
            'resources_affected': self.resources_affected or '',
            'evidence': self.evidence or '',
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M:%S')
        }

# User loader for Flask-Login
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Custom login_required decorator (optional, but can add extra checks)
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'admin':
            return jsonify({'error': 'Admin privileges required'}), 403
        return f(*args, **kwargs)
    return decorated_function

# Initialize database
def init_db():
    with app.app_context():
        db.create_all()
        
        # Create default admin user if it doesn't exist
        admin = User.query.filter_by(username='admin').first()
        if not admin:
            admin = User(
                username='admin',
                email='admin@example.com',
                role='admin'
            )
            admin.set_password('change_me')
            admin.set_keywords('security', 'pentest')
            db.session.add(admin)
            db.session.commit()
            print("Default admin user created. Please change the password immediately.")

# Routes for Web Pages
@app.route('/')
def index():
    return redirect(url_for('login_page'))

@app.route('/login')
def login_page():
    if current_user.is_authenticated:
        return redirect(url_for('app_page'))
    return render_template('login.html')

@app.route('/app')
@login_required
def app_page():
    return render_template('app.html')

# API Routes
@app.route('/api/check-auth', methods=['GET'])
def check_auth():
    if current_user.is_authenticated:
        return jsonify({
            'authenticated': True,
            'user': {
                'username': current_user.username,
                'role': current_user.role
            }
        })
    return jsonify({'authenticated': False}), 401

@app.route('/api/login', methods=['POST'])
def login():
    data = request.json
    user = User.query.filter_by(username=data['username']).first()
    
    if user and user.check_password(data['password']):
        # Log in the user
        login_user(user)
        
        user.last_login = datetime.utcnow()
        db.session.commit()
        
        return jsonify({
            'message': 'Login successful',
            'user': {
                'username': user.username,
                'role': user.role
            }
        })
    
    return jsonify({'error': 'Invalid username or password'}), 401

@app.route('/api/logout', methods=['POST'])
def logout():
    logout_user()
    return jsonify({'message': 'Logout successful'})

# User Findings Routes
@app.route('/api/user-findings', methods=['GET'])
@login_required
def get_user_findings():
    user_findings = UserFinding.query.filter_by(user_id=current_user.id).all()
    return jsonify([finding.to_dict() for finding in user_findings])

@app.route('/api/user-findings', methods=['POST'])
@login_required
def add_user_finding():
    # Verify request is JSON
    if not request.is_json:
        return jsonify({'error': 'Request must be JSON'}), 400
    
    # Parse incoming JSON data
    data = request.get_json()
    finding_id = data.get('finding_id')
    
    if not finding_id:
        return jsonify({'error': 'Finding ID is required'}), 400
    
    try:
        # Check if this finding is already added by this user
        existing = UserFinding.query.filter_by(
            user_id=current_user.id,
            finding_id=finding_id
        ).first()
        
        if existing:
            return jsonify({'error': 'Finding already added'}), 400
        
        # Get original finding data
        finding = Finding.query.get_or_404(finding_id)
        
        # Create user finding
        user_finding = UserFinding(
            user_id=current_user.id,
            finding_id=finding_id,
            title=finding.title,
            risk_rating=finding.risk_rating,
            description=finding.description,
            impact=finding.impact,
            resolution=finding.resolution,
            category=finding.category
        )
        
        db.session.add(user_finding)
        db.session.commit()
        
        return jsonify(user_finding.to_dict()), 201
    
    except Exception as e:
        # Log the full error for debugging
        print(f"Error adding user finding: {str(e)}")
        db.session.rollback()
        return jsonify({'error': 'Failed to add finding'}), 500

@app.route('/api/user-findings/<int:id>', methods=['DELETE'])
@login_required
def delete_user_finding(id):
    user_finding = UserFinding.query.filter_by(
        id=id, 
        user_id=current_user.id
    ).first_or_404()
    
    db.session.delete(user_finding)
    
    try:
        db.session.commit()
        return jsonify({'message': 'User finding deleted'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/user-findings/delete-all', methods=['POST'])
@login_required
def delete_all_user_findings():
    UserFinding.query.filter_by(user_id=current_user.id).delete()
    
    try:
        db.session.commit()
        return jsonify({'message': 'All user findings deleted'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# Findings Routes
@app.route('/api/findings', methods=['GET'])
@login_required
def get_findings():
    findings = Finding.query.all()
    return jsonify([finding.to_dict() for finding in findings])

@app.route('/api/findings', methods=['POST'])
@login_required
def add_finding():
    data = request.json
    new_finding = Finding(
        title=data['title'],
        risk_rating=data['risk_rating'],
        description=data['description'],
        impact=data['impact'],
        resolution=data['resolution'],
        category=data.get('category', '')
    )
    db.session.add(new_finding)
    db.session.commit()
    return jsonify(new_finding.to_dict()), 201

@app.route('/api/user-findings/<int:id>', methods=['PUT'])
@login_required
def update_user_finding(id):
    # Find the user finding that belongs to the current user
    user_finding = UserFinding.query.filter_by(
        id=id, 
        user_id=current_user.id
    ).first_or_404()
    
    # Verify request is JSON
    if not request.is_json:
        return jsonify({'error': 'Request must be JSON'}), 400
    
    # Parse incoming JSON data
    data = request.get_json()
    
    try:
        # Update fields that are allowed to be modified
        user_finding.title = data.get('title', user_finding.title)
        user_finding.risk_rating = data.get('risk_rating', user_finding.risk_rating)
        user_finding.description = data.get('description', user_finding.description)
        user_finding.impact = data.get('impact', user_finding.impact)
        user_finding.resolution = data.get('resolution', user_finding.resolution)
        user_finding.category = data.get('category', user_finding.category)
        
        # Optional fields for additional information
        user_finding.resources_affected = data.get('resources_affected', user_finding.resources_affected)
        user_finding.evidence = data.get('evidence', user_finding.evidence)
        
        db.session.commit()
        
        return jsonify(user_finding.to_dict()), 200
    
    except Exception as e:
        # Log the full error for debugging
        print(f"Error updating user finding: {str(e)}")
        db.session.rollback()
        return jsonify({'error': 'Failed to update finding'}), 500


@app.route('/api/findings/<int:id>', methods=['PUT'])
@login_required
def update_finding(id):
    finding = Finding.query.get_or_404(id)
    data = request.json
    finding.title = data['title']
    finding.risk_rating = data['risk_rating']
    finding.description = data['description']
    finding.impact = data['impact']
    finding.resolution = data['resolution']
    finding.category = data.get('category', '')
    db.session.commit()
    return jsonify(finding.to_dict())

@app.route('/api/findings/delete', methods=['POST'])
@login_required
def delete_findings():
    finding_ids = request.json.get('findings', [])
    Finding.query.filter(Finding.id.in_(finding_ids)).delete(synchronize_session=False)
    db.session.commit()
    return jsonify({'message': 'Findings deleted successfully'}), 200

@app.route('/api/export', methods=['POST'])
@login_required
def export_findings():
    data = request.json
    finding_ids = data.get('findings', [])
    resources = data.get('resources', {})
    evidence = data.get('evidence', {})
    edited_findings = data.get('edited_findings', {})

    # Get findings from the database
    db_findings = Finding.query.filter(Finding.id.in_(finding_ids)).all()
    findings_map = {str(f.id): f for f in db_findings}
    
    severity_order = {
        'Critical': 0,
        'High': 1,
        'Medium': 2,
        'Low': 3,
        'Informational': 4
    }
    
    findings_data = []
    
    for finding_id in finding_ids:
        str_id = str(finding_id)
        
        if str_id in edited_findings:
            finding_data = edited_findings[str_id]
            findings_data.append({
                'title': finding_data['title'],
                'severity': finding_data['risk_rating'],
                'description': finding_data['description'],
                'impact': finding_data['impact'],
                'resolution': finding_data['resolution'],
                'resources_affected': resources.get(str_id, ''),
                'evidence': evidence.get(str_id, '')
            })
        elif str_id in findings_map:
            finding = findings_map[str_id]
            findings_data.append({
                'title': finding.title,
                'severity': finding.risk_rating,
                'description': finding.description,
                'impact': finding.impact,
                'resolution': finding.resolution,
                'resources_affected': resources.get(str_id, ''),
                'evidence': evidence.get(str_id, '')
            })
    
    findings_data.sort(key=lambda x: severity_order.get(x['severity'], 999))
    context = {'findings': findings_data}
    
    doc_buffer = io.BytesIO()
    generate_docx(context, doc_buffer)
    
    return send_file(
        doc_buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='security_findings.docx'
    )

def generate_docx(context, buffer):
    """Generate DOCX file from context data."""
    doc = Document()
    for finding in context['findings']:
        doc.add_heading(finding['title'], level=1)
        doc.add_paragraph(f"Resources Affected: {finding['resources_affected']}")
        doc.add_paragraph(f"Severity: {finding['severity']}")
        doc.add_paragraph(f"Description: {finding['description']}")
        doc.add_paragraph(f"Impact: {finding['impact']}")
        doc.add_paragraph(f"Recommendations: {finding['resolution']}")
        doc.add_paragraph("References: ")
        doc.add_paragraph("Evidence and Reproduction Steps:")
        add_html_to_doc(doc, finding['evidence'])
        if finding != context['findings'][-1]:
            doc.add_page_break()
    doc.save(buffer)
    buffer.seek(0)


def add_html_to_doc(doc, html):
    """Convert HTML to docx content with proper list handling."""
    soup = BeautifulSoup(html, 'html.parser')
    list_counter = [0]  # Counter stack for numbered lists

    def process_element(element, parent=None, list_level=0, list_type=None):
        nonlocal list_counter
        
        if parent is None:
            parent = doc

        # Handle different element types
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            parent.add_heading(element.get_text(), level=level)
            return parent

        elif element.name == 'p':
            p = doc.add_paragraph()
            for child in element.contents:
                process_element(child, p, list_level, list_type)
            return p

        elif element.name in ['ul', 'ol']:
            # Reset counter for new list
            list_counter = list_counter[:list_level+1] + [0]
            new_list_type = 'bullet' if element.name == 'ul' else 'number'
            
            for child in element.find_all(['li'], recursive=False):
                process_element(child, parent, list_level + 1, new_list_type)
            return parent

        elif element.name == 'li':
            # Create appropriate list style
            if list_type == 'number':
                style = 'List Number'
                list_counter[list_level] += 1
                number = list_counter[list_level]
            else:
                style = 'List Bullet'
                number = None

            p = doc.add_paragraph(style=style)
            
            # Set indentation
            p.paragraph_format.left_indent = Inches(0.5 * list_level)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Add list content
            for child in element.contents:
                if child.name is None:  # Text node
                    text = str(child).strip()
                    if text:
                        if number and list_level == 0:
                            p.add_run(f"{number}. ").bold = True
                        p.add_run(text)
                else:
                    process_element(child, p, list_level, list_type)
            return parent

        elif element.name == 'br':
            if isinstance(parent, Paragraph):
                parent.add_run().add_break()
            return parent

        elif element.name in ['strong', 'b']:
            run = parent.add_run()
            run.bold = True
            run.text = element.get_text()
            return parent

        elif element.name in ['em', 'i']:
            run = parent.add_run()
            run.italic = True
            run.text = element.get_text()
            return parent

        elif element.name == 'u':
            run = parent.add_run()
            run.underline = True
            run.text = element.get_text()
            return parent

        # Handle text content
        if isinstance(element, str):
            text = element.strip()
            if text:
                if isinstance(parent, Paragraph):
                    parent.add_run(text)
                else:
                    doc.add_paragraph(text)
            return parent

        # Recursively process children
        for child in element.contents:
            process_element(child, parent, list_level, list_type)

        return parent

    
    for element in soup.contents:
        if element != '\n':
            process_element(element)
# Main execution
if __name__ == '__main__':
    init_db()
    app.run(debug=True)
